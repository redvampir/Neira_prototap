"""
Цель: Генерация ТКП в формате docx по каталогу и шаблону.
Инварианты: Шаблоны берём из training_data/tkp/samples, пропуски → "Требуется уточнение".
Риски: PDF извлекается шумно, отсутствует python-docx.
Проверка: python -m py_compile neira/organs/tkp/generator.py
"""

from __future__ import annotations

import difflib
import logging
import re
import zipfile
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Callable, Sequence

from neira.organs.tkp.parser import (
    ParsedCatalog,
    ParsedParameter,
    load_or_parse_catalog,
)

from neira.config import (
    TKP_CATALOGS_DIR,
    TKP_COMPANY_PARAGRAPH_LIMIT,
    TKP_INTRO_PARAGRAPH_LIMIT,
    TKP_MISSING_VALUE_TEXT,
    TKP_OUTPUT_DIR,
    TKP_PARSED_DIR,
    TKP_SAMPLES_DIR,
    TKP_PARAM_MATCH_MIN_SCORE,
)

logger = logging.getLogger(__name__)

_DOCX_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
_TEMPLATE_CACHE: dict[Path, "TkpTemplateSpec"] = {}
_RAW_PARAM_ALIASES: dict[str, list[str]] = {
    "Размер рабочего стола": [
        "table size",
        "table dimension",
        "table diameter",
        "table dia",
        "worktable size",
    ],
    "Максимальная нагрузка на стол": [
        "max table load",
        "maximum table load",
        "table load",
        "max load",
        "maximum load",
        "maximum workpiece weight",
        "max workpiece weight",
        "workpiece weight",
        "table workpiece weight",
        "max workpiece",
    ],
    "Угол наклона оси": [
        "tilt angle",
        "axis tilt",
        "tilting angle",
        "a axis",
        "a-axis",
    ],
    "Угол поворота": [
        "rotation angle",
        "rotary angle",
        "swivel angle",
        "c axis",
        "c-axis",
    ],
    "Вращение стола": [
        "table rotation",
        "table speed",
        "rotary table speed",
        "table rpm",
        "c axis speed",
        "caxis",
        "c axis",
        "rotating axis",
        "rapid rotating axis",
    ],
    "Перемещение по осям X": [
        "x-axis travel",
        "x travel",
        "travel x",
        "x-axis stroke",
        "travels x axis",
        "travel distance x axis",
    ],
    "Перемещение по оси X": [
        "x-axis travel",
        "x travel",
        "travel x",
        "x-axis stroke",
        "travels x axis",
        "travel distance x axis",
    ],
    "Перемещение по осям Y": [
        "y-axis travel",
        "y travel",
        "travel y",
        "y-axis stroke",
        "travels y axis",
        "travel distance y axis",
    ],
    "Перемещение по оси Y": [
        "y-axis travel",
        "y travel",
        "travel y",
        "y-axis stroke",
        "travels y axis",
        "travel distance y axis",
    ],
    "Перемещение по осям Z": [
        "z-axis travel",
        "z travel",
        "travel z",
        "z-axis stroke",
        "travels z axis",
        "travel distance z axis",
    ],
    "Перемещение по оси Z": [
        "z-axis travel",
        "z travel",
        "travel z",
        "z-axis stroke",
        "travels z axis",
        "travel distance z axis",
    ],
    "Зона обработки": [
        "machining area",
        "work area",
        "working area",
        "work zone",
        "workspace",
        "maximum workpiece size",
        "max workpiece size",
        "workpiece size",
    ],
    "Тип конического отверстия": [
        "spindle taper",
        "taper type",
        "taper hole",
        "tool shank",
        "tool shank type",
        "tool type",
    ],
    "Скорость": [
        "spindle speed",
        "max speed",
        "speed",
        "rpm",
    ],
    "Мощность": [
        "power",
        "spindle power",
        "kw",
    ],
    "Крутящий момент": [
        "torque",
        "max torque",
        "nm",
    ],
    "Быстрый ход по оси X": [
        "rapid traverse x",
        "rapid feed x",
        "rapid travel x",
    ],
    "Быстрый ход по оси Y": [
        "rapid traverse y",
        "rapid feed y",
        "rapid travel y",
    ],
    "Быстрый ход по оси Z": [
        "rapid traverse z",
        "rapid feed z",
        "rapid travel z",
    ],
    "Максимальная скорость вращения оси В": [
        "b-axis max speed",
        "b axis speed",
        "b-axis speed",
        "baxis",
        "tilting axis speed",
    ],
    "Максимальная скорость вращения оси C": [
        "c-axis max speed",
        "c axis speed",
        "caxis",
        "rotating axis",
    ],
    "Головка шпинделя": [
        "spindle head",
        "spindle nose",
        "spindle type",
    ],
    "Макс.скорость шпинделя": [
        "max spindle speed",
        "spindle speed",
    ],
    "Мощность шпинделя": [
        "spindle power",
        "motor power",
    ],
    "Макс. крутящий момент": [
        "max torque",
        "torque",
    ],
    "Диаметр патрона": [
        "chuck diameter",
        "chuck size",
    ],
    "Диаметр отверстия шпинделя": [
        "spindle bore",
        "spindle hole",
        "bore diameter",
    ],
    "Макс. диаметр стержня": [
        "max bar diameter",
        "bar capacity",
    ],
    "Смазка подшипника шпинделя": [
        "spindle bearing lubrication",
        "bearing lubrication",
    ],
    "Смазка направляющей": [
        "guideway lubrication",
        "way lubrication",
    ],
    "Максимальное вращение над станиной": [
        "swing over bed",
        "max swing over bed",
    ],
    "Максимальное вращение над суппортом": [
        "swing over cross slide",
        "max swing over cross slide",
    ],
    "Максимальный диаметр обработки": [
        "max turning diameter",
        "max machining diameter",
    ],
    "Рекомендуемый макс. диаметр обработки": [
        "recommended max turning diameter",
    ],
    "Быстрое перемещение по оси Z": [
        "rapid traverse z",
        "rapid feed z",
    ],
    "Быстрое перемещение по оси X": [
        "rapid traverse x",
        "rapid feed x",
    ],
    "Макс. Точность позиционирования (X / Z); VDI 3441": [
        "positioning accuracy",
        "vdi 3441 positioning",
    ],
    "Точность повторяемости (X / Z); VDI 3441": [
        "repeatability",
        "repeat accuracy",
        "vdi 3441 repeatability",
    ],
    "Револьверная головка": [
        "turret",
        "tool turret",
        "revolver",
    ],
}


class TkpGenerationError(RuntimeError):
    """Ошибки генерации ТКП."""


@dataclass(frozen=True)
class TkpTechRow:
    group: str | None
    name: str
    unit: str | None
    is_group: bool = False


@dataclass(frozen=True)
class TkpCommercialTerm:
    key: str
    value: str | None


@dataclass(frozen=True)
class TkpTemplateSpec:
    template_id: str
    sample_path: Path
    sample_model: str | None
    intro_paragraphs: list[str]
    company_paragraphs: list[str]
    commercial_terms: list[TkpCommercialTerm]
    tech_rows: list[TkpTechRow]
    tech_has_group_column: bool
    price_headers: list[str]
    options_headers: list[str]
    general_info_headers: list[str] | None
    general_info_rows: list[str]


@dataclass(frozen=True)
class TkpCatalogInfo:
    catalog_path: Path
    model_snippet: str | None
    main_units: list[str]
    standard_items: list[str]
    option_items: list[str]
    param_values: dict[str, str]


@dataclass(frozen=True)
class TkpGenerationResult:
    output_path: Path
    template_id: str
    catalog_path: Path
    warnings: list[str]
    missing_values: int


def generate_tkp_document(
    model_name: str,
    catalogs_dir: Path | None = None,
    samples_dir: Path | None = None,
    output_dir: Path | None = None,
    parsed_dir: Path | None = None,
    missing_value_text: str | None = None,
    force_parse: bool = False,
) -> TkpGenerationResult:
    """
    Сгенерировать docx ТКП по модели.

    Args:
        model_name: Имя модели станка (например, "DVF 6500T").
        catalogs_dir: Папка с PDF каталогами.
        samples_dir: Папка с docx шаблонами.
        output_dir: Папка для сохранения docx.
        parsed_dir: Папка для JSON/MD парсинга.
        missing_value_text: Текст для пропусков.
        force_parse: Принудительно пересобрать JSON/MD.
    """
    model_name = _normalize_model_input(model_name)
    catalogs_dir, samples_dir, output_dir, parsed_dir, missing_value_text = (
        _resolve_generation_paths(
            catalogs_dir,
            samples_dir,
            output_dir,
            parsed_dir,
            missing_value_text,
        )
    )
    _validate_generation_dirs(catalogs_dir, samples_dir)

    template, catalog_path, warnings, catalog_info = _prepare_generation_context(
        model_name=model_name,
        catalogs_dir=catalogs_dir,
        samples_dir=samples_dir,
        parsed_dir=parsed_dir,
        force_parse=force_parse,
    )

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = _build_output_path(output_dir, model_name)

    missing_values = _render_docx(
        output_path=output_path,
        model_name=model_name,
        template=template,
        catalog_info=catalog_info,
        missing_value_text=missing_value_text,
    )

    return TkpGenerationResult(
        output_path=output_path,
        template_id=template.template_id,
        catalog_path=catalog_path,
        warnings=warnings,
        missing_values=missing_values,
    )


def _normalize_model_input(model_name: str) -> str:
    model_name = model_name.strip()
    if not model_name:
        raise TkpGenerationError("Не указана модель станка.")
    return model_name


def _resolve_generation_paths(
    catalogs_dir: Path | None,
    samples_dir: Path | None,
    output_dir: Path | None,
    parsed_dir: Path | None,
    missing_value_text: str | None,
) -> tuple[Path, Path, Path, Path, str]:
    return (
        catalogs_dir or TKP_CATALOGS_DIR,
        samples_dir or TKP_SAMPLES_DIR,
        output_dir or TKP_OUTPUT_DIR,
        parsed_dir or TKP_PARSED_DIR,
        missing_value_text or TKP_MISSING_VALUE_TEXT,
    )


def _validate_generation_dirs(catalogs_dir: Path, samples_dir: Path) -> None:
    if not catalogs_dir.exists():
        raise TkpGenerationError(f"Каталоги не найдены: {catalogs_dir}")
    if not samples_dir.exists():
        raise TkpGenerationError(f"Шаблоны не найдены: {samples_dir}")


def _prepare_generation_context(
    model_name: str,
    catalogs_dir: Path,
    samples_dir: Path,
    parsed_dir: Path,
    force_parse: bool,
) -> tuple[TkpTemplateSpec, Path, list[str], TkpCatalogInfo]:
    templates = _load_templates(samples_dir)
    catalog_path, warnings = _select_catalog(catalogs_dir, model_name)
    template = _select_template(templates, catalog_path, model_name)

    parsed_result = load_or_parse_catalog(
        catalog_path=catalog_path,
        model_name=model_name,
        output_dir=parsed_dir,
        force_parse=force_parse,
    )
    warnings.extend(parsed_result.warnings)
    catalog_info = _extract_catalog_info(parsed_result.parsed, template)
    return template, catalog_path, warnings, catalog_info


def _load_templates(samples_dir: Path) -> list[TkpTemplateSpec]:
    samples = sorted(samples_dir.glob("*.docx"))
    if not samples:
        raise TkpGenerationError(f"В папке нет docx-шаблонов: {samples_dir}")

    templates: list[TkpTemplateSpec] = []
    for sample_path in samples:
        templates.append(_build_template_spec(sample_path))
    return templates


def _build_template_spec(sample_path: Path) -> TkpTemplateSpec:
    if sample_path in _TEMPLATE_CACHE:
        return _TEMPLATE_CACHE[sample_path]

    paragraphs = _extract_docx_paragraphs(sample_path)
    tables = _extract_docx_tables(sample_path)

    intro_paragraphs = _pick_intro_paragraphs(paragraphs)
    company_paragraphs = _pick_company_paragraphs(paragraphs)

    commercial_table = _find_table(tables, _is_commercial_table)
    commercial_terms = _parse_commercial_terms(commercial_table)

    tech_table = _find_table(tables, _is_tech_table)
    tech_rows, tech_has_group, header_model = _parse_tech_table(tech_table)

    price_table = _find_table(tables, _is_price_table)
    options_table = _find_table(tables, _is_options_table)
    general_info_table = _find_table(tables, _is_general_info_table)

    price_headers = price_table[0] if price_table else ["No.", "Наименование", "Цена"]
    options_headers = options_table[0] if options_table else ["No.", "Опции", "Цена"]
    general_info_headers = general_info_table[0] if general_info_table else None
    general_info_rows = _parse_general_info_rows(general_info_table)

    sample_model = header_model or _extract_model_from_paragraphs(paragraphs)
    template_id = _infer_template_id(sample_path)

    spec = TkpTemplateSpec(
        template_id=template_id,
        sample_path=sample_path,
        sample_model=sample_model,
        intro_paragraphs=intro_paragraphs,
        company_paragraphs=company_paragraphs,
        commercial_terms=commercial_terms,
        tech_rows=tech_rows,
        tech_has_group_column=tech_has_group,
        price_headers=price_headers,
        options_headers=options_headers,
        general_info_headers=general_info_headers,
        general_info_rows=general_info_rows,
    )
    _TEMPLATE_CACHE[sample_path] = spec
    return spec


def _extract_docx_paragraphs(sample_path: Path) -> list[str]:
    with zipfile.ZipFile(sample_path) as archive:
        xml_data = archive.read("word/document.xml")
    root = ET.fromstring(xml_data)

    paragraphs: list[str] = []
    for p in root.findall(".//w:p", _DOCX_NS):
        texts = [t.text for t in p.findall(".//w:t", _DOCX_NS) if t.text]
        if texts:
            paragraph = "".join(texts).strip()
            if paragraph:
                paragraphs.append(paragraph)
    return paragraphs


def _extract_docx_tables(sample_path: Path) -> list[list[list[str]]]:
    with zipfile.ZipFile(sample_path) as archive:
        xml_data = archive.read("word/document.xml")
    root = ET.fromstring(xml_data)

    tables: list[list[list[str]]] = []
    for tbl in root.findall(".//w:tbl", _DOCX_NS):
        rows: list[list[str]] = []
        for tr in tbl.findall("./w:tr", _DOCX_NS):
            cells: list[str] = []
            for tc in tr.findall("./w:tc", _DOCX_NS):
                texts = [t.text for t in tc.findall(".//w:t", _DOCX_NS) if t.text]
                cells.append("".join(texts).strip())
            rows.append(cells)
        tables.append(rows)
    return tables


def _pick_intro_paragraphs(paragraphs: Sequence[str]) -> list[str]:
    intro: list[str] = []
    for paragraph in paragraphs:
        lowered = paragraph.lower()
        if "гарант" in lowered or "пуско-наладоч" in lowered:
            continue
        if len(intro) < TKP_INTRO_PARAGRAPH_LIMIT:
            intro.append(paragraph)
        if len(intro) >= TKP_INTRO_PARAGRAPH_LIMIT:
            break
    return intro


def _pick_company_paragraphs(paragraphs: Sequence[str]) -> list[str]:
    idx = next(
        (i for i, p in enumerate(paragraphs) if "о компании" in p.lower()),
        None,
    )
    if idx is None:
        return []

    collected: list[str] = []
    for paragraph in paragraphs[idx + 1 :]:
        lowered = paragraph.lower()
        if any(
            marker in lowered
            for marker in ("технические параметры", "описание серии", "описание станков")
        ):
            break
        if paragraph.strip():
            collected.append(paragraph.strip())
        if len(collected) >= TKP_COMPANY_PARAGRAPH_LIMIT:
            break
    return collected


def _find_table(
    tables: Sequence[list[list[str]]],
    predicate: Callable[[list[list[str]]], bool],
) -> list[list[str]]:
    for table in tables:
        if predicate(table):
            return table
    return []


def _is_commercial_table(table: list[list[str]]) -> bool:
    if not table or not table[0]:
        return False
    if len(table[0]) != 2:
        return False
    return any("гарант" in cell.lower() for cell in table[0] if cell)


def _is_tech_table(table: list[list[str]]) -> bool:
    if not table or not table[0]:
        return False
    header = " ".join(cell.lower() for cell in table[0] if cell)
    return "параметр" in header or "технические" in header


def _is_price_table(table: list[list[str]]) -> bool:
    if not table or not table[0]:
        return False
    header = " ".join(cell.lower() for cell in table[0] if cell)
    return "цена" in header and "наимен" in header and "опци" not in header


def _is_options_table(table: list[list[str]]) -> bool:
    if not table or not table[0]:
        return False
    header = " ".join(cell.lower() for cell in table[0] if cell)
    return "опци" in header


def _is_general_info_table(table: list[list[str]]) -> bool:
    if not table or not table[0]:
        return False
    header = " ".join(cell.lower() for cell in table[0] if cell)
    return "производител" in header and "страна" in header


def _parse_commercial_terms(table: list[list[str]]) -> list[TkpCommercialTerm]:
    terms: list[TkpCommercialTerm] = []
    for row in table:
        if len(row) < 2:
            continue
        key = row[0].strip()
        value = row[1].strip() if len(row) > 1 else ""
        if not key:
            continue
        if "срок изготовления" in key.lower() and "срок доставки" in key.lower():
            terms.append(TkpCommercialTerm("Срок изготовления", None))
            terms.append(TkpCommercialTerm("Срок доставки", None))
            continue
        terms.append(TkpCommercialTerm(key, value or None))
    return terms


def _parse_general_info_rows(table: list[list[str]]) -> list[str]:
    if not table:
        return []
    rows: list[str] = []
    for row in table[1:]:
        if row:
            label = row[0].strip()
            if label:
                rows.append(label)
    return rows


def _parse_tech_table(
    table: list[list[str]],
) -> tuple[list[TkpTechRow], bool, str | None]:
    if not table:
        return [], False, None

    header = table[0]
    header_model = _extract_model_from_header(header)
    has_group_column = any("ед." in cell.lower() for cell in header if cell)

    rows: list[TkpTechRow] = []
    current_group: str | None = None

    for row in table[1:]:
        if not row:
            continue
        if len(row) == 1:
            current_group = row[0].strip() or current_group
            if current_group:
                rows.append(TkpTechRow(current_group, current_group, None, is_group=True))
            continue

        if has_group_column:
            group = row[0].strip() or current_group
            name = row[1].strip() if len(row) > 1 else ""
            unit = row[3].strip() if len(row) > 3 else ""
            current_group = group or current_group
            if name:
                rows.append(TkpTechRow(group, name, unit or None))
        else:
            name = row[0].strip()
            unit = row[1].strip() if len(row) > 1 else ""
            if name:
                rows.append(TkpTechRow(current_group, name, unit or None))

    return rows, has_group_column, header_model


def _extract_model_from_header(header: Sequence[str]) -> str | None:
    candidates = [
        cell.strip()
        for cell in header
        if cell and "параметр" not in cell.lower() and "ед." not in cell.lower()
    ]
    return candidates[0] if candidates else None


def _extract_model_from_paragraphs(paragraphs: Sequence[str]) -> str | None:
    for paragraph in paragraphs:
        match = re.search(r"модель\\s+([A-Za-zА-Яа-я0-9-]+)", paragraph, re.IGNORECASE)
        if match:
            return match.group(1)
    for paragraph in paragraphs:
        match = re.search(r"\\b([A-Z]{1,4}\\s?\\d{2,4}[A-Z]?)\\b", paragraph)
        if match:
            return match.group(1).strip()
    return None


def _infer_template_id(sample_path: Path) -> str:
    name = sample_path.stem.lower()
    if "dvf" in name:
        return "dvf"
    if "ht" in name or "ct" in name:
        return "ht"
    return re.sub(r"[^a-z0-9]+", "_", name).strip("_") or "tkp"


def _select_catalog(
    catalogs_dir: Path,
    model_name: str,
) -> tuple[Path, list[str]]:
    catalogs = sorted(catalogs_dir.glob("*.pdf"))
    if not catalogs:
        raise TkpGenerationError(f"В каталоге нет PDF: {catalogs_dir}")

    normalized_model = _normalize_model(model_name)
    warnings: list[str] = []

    prefix = _extract_model_prefix(model_name)
    prefix_match = _select_catalog_by_prefix(catalogs, prefix)
    if prefix_match is not None:
        return prefix_match, warnings

    if normalized_model:
        for catalog in catalogs:
            if normalized_model in _normalize_model(catalog.stem):
                return catalog, warnings

    best_path = catalogs[0]
    best_score = 0
    for catalog in catalogs:
        text = _load_catalog_text(catalog)
        score = _score_model_in_text(text, normalized_model)
        if score > best_score:
            best_score = score
            best_path = catalog

    if best_score == 0:
        warnings.append(
            "Модель не найдена в каталогах. Проверь название - выбран первый каталог."
        )

    return best_path, warnings


def _select_template(
    templates: Sequence[TkpTemplateSpec],
    catalog_path: Path,
    model_name: str,
) -> TkpTemplateSpec:
    name = catalog_path.name.lower()
    model = model_name.lower()

    template = _find_template_by_hint(templates, name)
    if template:
        return template

    template = _find_template_by_hint(templates, model)
    if template:
        return template

    return templates[0]


def _find_template_by_hint(
    templates: Sequence[TkpTemplateSpec],
    hint: str,
) -> TkpTemplateSpec | None:
    if "dvf" in hint:
        return next((t for t in templates if t.template_id == "dvf"), None)
    if "ht" in hint or "ct" in hint:
        return next((t for t in templates if t.template_id == "ht"), None)
    return None


def _extract_model_prefix(model_name: str) -> str:
    match = re.match(r"[A-Za-z]+", model_name.strip())
    return match.group(0).lower() if match else ""


def _select_catalog_by_prefix(
    catalogs: Sequence[Path],
    prefix: str,
) -> Path | None:
    if not prefix:
        return None
    prefix_lower = prefix.lower()
    if prefix_lower.startswith("dvf"):
        return _find_catalog_by_keywords(catalogs, ("dvf",))
    if prefix_lower.startswith(("ht", "ct")):
        return _find_catalog_by_keywords(catalogs, ("turning", "ht", "ct"))
    return None


def _find_catalog_by_keywords(
    catalogs: Sequence[Path],
    keywords: Sequence[str],
) -> Path | None:
    for catalog in catalogs:
        name = catalog.stem.lower()
        if any(keyword in name for keyword in keywords):
            return catalog
    return None


def _load_catalog_text(path: Path) -> str:
    try:
        import pypdf
    except ImportError as exc:
        raise TkpGenerationError(
            "Для чтения PDF нужен пакет pypdf: pip install pypdf"
        ) from exc

    reader = pypdf.PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text:
            parts.append(text)
    return "\n".join(parts)


def _normalize_model(text: str) -> str:
    return re.sub(r"[^a-zA-Z0-9]+", "", text).lower()


def _score_model_in_text(text: str, normalized_model: str) -> int:
    if not text or not normalized_model:
        return 0
    normalized_text = _normalize_model(text)
    return normalized_text.count(normalized_model)


def _extract_catalog_info(
    parsed: ParsedCatalog,
    template: TkpTemplateSpec,
) -> TkpCatalogInfo:
    param_values = _map_param_values(parsed.parameters, template.tech_rows)
    return TkpCatalogInfo(
        catalog_path=Path(parsed.catalog_path),
        model_snippet=parsed.snippet,
        main_units=parsed.main_units,
        standard_items=parsed.standard_items,
        option_items=parsed.option_items,
        param_values=param_values,
    )


def _map_param_values(
    parameters: Sequence[ParsedParameter],
    tech_rows: Sequence[TkpTechRow],
) -> dict[str, str]:
    indexed = _index_parameters(parameters)
    alias_map = _build_param_aliases()
    values: dict[str, str] = {}

    for row in tech_rows:
        if row.is_group:
            continue
        matched = _match_param(row.name, indexed, alias_map, row.unit)
        if matched:
            values[row.name] = matched.value
    _augment_param_values(values, parameters, tech_rows)
    return values


def _index_parameters(
    parameters: Sequence[ParsedParameter],
) -> dict[str, ParsedParameter]:
    indexed: dict[str, ParsedParameter] = {}
    for param in parameters:
        key = _normalize_param_name(param.name)
        if not key:
            continue
        existing = indexed.get(key)
        if existing is None or len(param.value) > len(existing.value):
            indexed[key] = param
    return indexed


def _augment_param_values(
    values: dict[str, str],
    parameters: Sequence[ParsedParameter],
    tech_rows: Sequence[TkpTechRow],
) -> None:
    controllers = _extract_controllers(parameters)
    dimensions = _extract_dimensions(parameters)
    weight_kg = _extract_weight_kg(parameters)

    for row in tech_rows:
        if row.is_group or row.name in values:
            continue
        group = (row.group or "").lower()
        if "чпу" in group and controllers:
            values[row.name] = controllers
        elif "габарит" in group and dimensions:
            values[row.name] = dimensions
        elif "вес" in group and weight_kg is not None:
            values[row.name] = _format_weight_value(weight_kg, row.unit)


def _extract_controllers(parameters: Sequence[ParsedParameter]) -> str | None:
    tokens = ("fanuc", "siemens", "heidenhain")
    normalized = [_normalize_param_name(param.name) for param in parameters]
    found: list[str] = []
    for token in tokens:
        if any(token in name for name in normalized):
            found.append(token.upper())
    return " / ".join(found) if found else None


def _extract_dimensions(parameters: Sequence[ParsedParameter]) -> str | None:
    dims: dict[str, float] = {}
    for param in parameters:
        key = _normalize_param_name(param.name)
        if "dimension" not in key:
            continue
        if "length" in key:
            value = _parse_numeric_value(param.value)
            if value is not None:
                dims["length"] = value
        elif "width" in key:
            value = _parse_numeric_value(param.value)
            if value is not None:
                dims["width"] = value
        elif "height" in key:
            value = _parse_numeric_value(param.value)
            if value is not None:
                dims["height"] = value

    ordered = ["length", "width", "height"]
    parts = [_format_number(dims[item]) for item in ordered if item in dims]
    if len(parts) < 2:
        return None
    return " x ".join(parts)


def _extract_weight_kg(parameters: Sequence[ParsedParameter]) -> float | None:
    for param in parameters:
        key = _normalize_param_name(param.name)
        if "dimension" in key and "weight" in key:
            return _parse_numeric_value(param.value)
    return None


def _parse_numeric_value(value: str) -> float | None:
    match = re.search(r"[-+]?\\d+(?:[\\.,]\\d+)?", value)
    if not match:
        return None
    raw = match.group(0).replace(" ", "").replace(",", ".")
    try:
        return float(raw)
    except ValueError:
        return None


def _format_weight_value(weight_kg: float, unit: str | None) -> str:
    if unit and unit.strip().lower().startswith(("t", "т")):
        return _format_number(weight_kg / 1000.0)
    return _format_number(weight_kg)


def _format_number(value: float) -> str:
    if float(value).is_integer():
        return str(int(value))
    return f"{value:.1f}".rstrip("0").rstrip(".")


def _match_param(
    name: str,
    indexed: dict[str, ParsedParameter],
    alias_map: dict[str, list[str]],
    unit: str | None,
) -> ParsedParameter | None:
    normalized = _normalize_param_name(name)
    if normalized in indexed:
        return indexed[normalized]

    aliases = alias_map.get(normalized, [])
    alias_match = _match_by_aliases(indexed, aliases, unit)
    if alias_match:
        return alias_match

    return _match_fuzzy(indexed, normalized, unit)


def _match_by_aliases(
    indexed: dict[str, ParsedParameter],
    aliases: Sequence[str],
    unit: str | None,
) -> ParsedParameter | None:
    if not aliases:
        return None
    best_param: ParsedParameter | None = None
    best_score = -1.0
    for alias in aliases:
        for key, param in indexed.items():
            if alias in key or key in alias:
                score = float(len(key))
                if _units_match(unit, param.unit):
                    score += 1000.0
                if score > best_score:
                    best_score = score
                    best_param = param
    return best_param


def _match_fuzzy(
    indexed: dict[str, ParsedParameter],
    normalized: str,
    unit: str | None,
) -> ParsedParameter | None:
    if not normalized:
        return None
    best_param: ParsedParameter | None = None
    best_score = 0.0
    for key, param in indexed.items():
        score = difflib.SequenceMatcher(None, normalized, key).ratio()
        if _units_match(unit, param.unit):
            score += 1.0
        if score > best_score:
            best_score = score
            best_param = param
    if best_score >= TKP_PARAM_MATCH_MIN_SCORE:
        return best_param
    return None


def _normalize_param_name(text: str) -> str:
    return re.sub(r"[^a-z0-9а-яё]+", " ", text.lower()).strip()


def _normalize_unit(unit: str) -> str:
    normalized = unit.lower().strip()
    normalized = normalized.replace(" ", "")
    normalized = normalized.replace("°", "deg")
    normalized = normalized.replace("r/min", "rpm")
    normalized = normalized.replace("r\\min", "rpm")
    normalized = normalized.replace("об/мин", "rpm")
    normalized = normalized.replace("квт", "kw")
    normalized = normalized.replace("нм", "nm")
    normalized = normalized.replace("кг", "kg")
    return normalized


def _units_match(row_unit: str | None, param_unit: str | None) -> bool:
    if not row_unit or not param_unit:
        return False
    row_norm = _normalize_unit(row_unit)
    param_norm = _normalize_unit(param_unit)
    if not row_norm or not param_norm:
        return False
    return row_norm in param_norm or param_norm in row_norm


def _build_param_aliases() -> dict[str, list[str]]:
    normalized: dict[str, list[str]] = {}
    for key, aliases in _RAW_PARAM_ALIASES.items():
        normalized_key = _normalize_param_name(key)
        normalized[normalized_key] = [_normalize_param_name(alias) for alias in aliases]
    return normalized


def _build_output_path(output_dir: Path, model_name: str) -> Path:
    safe_model = re.sub(r"[^a-zA-Z0-9._-]+", "_", model_name).strip("_")
    if not safe_model:
        safe_model = "model"
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return output_dir / f"TKP_{safe_model}_{timestamp}.docx"


def _render_docx(
    output_path: Path,
    model_name: str,
    template: TkpTemplateSpec,
    catalog_info: TkpCatalogInfo,
    missing_value_text: str,
) -> int:
    Document = _require_docx()
    document = Document()

    _render_intro_section(document, template, model_name, missing_value_text)
    _render_company_section(document, template, missing_value_text)

    missing_values = 0
    missing_values += _render_machine_info(document, catalog_info, missing_value_text)
    missing_values += _render_main_units(document, template, catalog_info, missing_value_text)
    missing_values += _render_tech_specs(document, template, catalog_info, missing_value_text)
    missing_values += _render_standard_equipment(document, catalog_info, missing_value_text)
    missing_values += _render_options(document, template, catalog_info, missing_value_text)
    _render_price_list(document, template)

    document.save(output_path)
    logger.info("Сформирован ТКП: %s", output_path)
    return missing_values


def _require_docx():
    try:
        from docx import Document
    except ImportError as exc:
        raise TkpGenerationError(
            "Для генерации docx установи python-docx: pip install python-docx"
        ) from exc
    return Document


def _apply_model_replacement(
    paragraphs: Sequence[str],
    sample_model: str | None,
    target_model: str,
) -> list[str]:
    if not sample_model:
        return list(paragraphs)
    pattern = re.compile(re.escape(sample_model), re.IGNORECASE)
    return [pattern.sub(target_model, paragraph) for paragraph in paragraphs]


def _render_intro_section(
    document,
    template: TkpTemplateSpec,
    model_name: str,
    missing_value_text: str,
) -> None:
    title = "Коммерческое предложение"
    document.add_heading(title, level=1)

    intro_paragraphs = _apply_model_replacement(
        template.intro_paragraphs, template.sample_model, model_name
    )
    for paragraph in intro_paragraphs:
        document.add_paragraph(paragraph)

    if template.commercial_terms:
        _add_commercial_terms_table(
            document,
            template.commercial_terms,
            missing_value_text,
        )


def _render_company_section(
    document,
    template: TkpTemplateSpec,
    missing_value_text: str,
) -> None:
    document.add_heading("Введение о компании", level=2)
    company_text = template.company_paragraphs or [missing_value_text]
    for paragraph in company_text:
        document.add_paragraph(paragraph)


def _render_machine_info(
    document,
    catalog_info: TkpCatalogInfo,
    missing_value_text: str,
) -> int:
    document.add_heading("Общая информация о станке", level=2)
    document.add_paragraph(catalog_info.model_snippet or missing_value_text)
    return 1 if catalog_info.model_snippet is None else 0


def _render_main_units(
    document,
    template: TkpTemplateSpec,
    catalog_info: TkpCatalogInfo,
    missing_value_text: str,
) -> int:
    document.add_heading("Информация по основным узлам", level=2)
    if template.general_info_headers:
        return _add_general_info_table(
            document,
            template,
            missing_value_text,
        )
    return _add_bullet_list(
        document,
        catalog_info.main_units,
        missing_value_text,
    )


def _render_tech_specs(
    document,
    template: TkpTemplateSpec,
    catalog_info: TkpCatalogInfo,
    missing_value_text: str,
) -> int:
    document.add_heading("Технические характеристики", level=2)
    return _add_tech_table(
        document,
        template,
        catalog_info.param_values,
        missing_value_text,
    )


def _render_standard_equipment(
    document,
    catalog_info: TkpCatalogInfo,
    missing_value_text: str,
) -> int:
    document.add_heading("Стандартная комплектация", level=2)
    return _add_bullet_list(
        document,
        catalog_info.standard_items,
        missing_value_text,
    )


def _render_options(
    document,
    template: TkpTemplateSpec,
    catalog_info: TkpCatalogInfo,
    missing_value_text: str,
) -> int:
    document.add_heading("Возможные опции", level=2)
    if template.options_headers:
        return _add_options_table(
            document,
            template.options_headers,
            catalog_info.option_items,
            missing_value_text,
        )
    return _add_bullet_list(
        document,
        catalog_info.option_items,
        missing_value_text,
    )


def _render_price_list(
    document,
    template: TkpTemplateSpec,
) -> None:
    document.add_heading("Прайс-лист", level=2)
    _add_empty_table(document, template.price_headers)


def _add_commercial_terms_table(
    document,
    terms: Sequence[TkpCommercialTerm],
    missing_value_text: str,
) -> None:
    table = document.add_table(rows=1, cols=2)
    header = table.rows[0].cells
    header[0].text = "Пункт"
    header[1].text = "Значение"

    for term in terms:
        row_cells = table.add_row().cells
        row_cells[0].text = term.key
        row_cells[1].text = term.value or missing_value_text


def _add_general_info_table(
    document,
    template: TkpTemplateSpec,
    missing_value_text: str,
) -> int:
    if not template.general_info_headers:
        return 0
    table = document.add_table(rows=1, cols=len(template.general_info_headers))
    header_cells = table.rows[0].cells
    for idx, header in enumerate(template.general_info_headers):
        header_cells[idx].text = header

    missing_values = 0
    if not template.general_info_rows:
        row_cells = table.add_row().cells
        row_cells[0].text = missing_value_text
        for cell in row_cells[1:]:
            cell.text = missing_value_text
        return len(row_cells)

    for row_label in template.general_info_rows:
        row_cells = table.add_row().cells
        row_cells[0].text = row_label
        for cell in row_cells[1:]:
            cell.text = missing_value_text
            missing_values += 1
    return missing_values


def _add_tech_table(
    document,
    template: TkpTemplateSpec,
    param_values: dict[str, str],
    missing_value_text: str,
) -> int:
    if template.tech_has_group_column:
        columns = 4
    else:
        columns = 3

    table = document.add_table(rows=1, cols=columns)
    header_cells = table.rows[0].cells

    if template.tech_has_group_column:
        header_cells[0].text = "Раздел"
        header_cells[1].text = "Параметр"
        header_cells[2].text = "Значение"
        header_cells[3].text = "Ед. изм."
    else:
        header_cells[0].text = "Параметр"
        header_cells[1].text = "Ед. изм."
        header_cells[2].text = "Значение"

    missing_values = 0
    last_group: str | None = None

    for row in template.tech_rows:
        row_cells = table.add_row().cells
        if row.is_group and not template.tech_has_group_column:
            row_cells[0].text = row.name
            row_cells[0].merge(row_cells[-1])
            continue

        value = param_values.get(row.name, missing_value_text)
        if value == missing_value_text:
            missing_values += 1

        if template.tech_has_group_column:
            group_value = row.group or ""
            if group_value == last_group:
                group_value = ""
            row_cells[0].text = group_value
            row_cells[1].text = row.name
            row_cells[2].text = value
            row_cells[3].text = row.unit or ""
            last_group = row.group or last_group
        else:
            row_cells[0].text = row.name
            row_cells[1].text = row.unit or ""
            row_cells[2].text = value

    return missing_values


def _add_bullet_list(
    document,
    items: Sequence[str],
    missing_value_text: str,
) -> int:
    if not items:
        document.add_paragraph(missing_value_text)
        return 1

    for item in items:
        paragraph = document.add_paragraph(item)
        try:
            paragraph.style = "List Bullet"
        except (KeyError, ValueError):
            pass
    return 0


def _add_empty_table(
    document,
    headers: Sequence[str],
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
    table.add_row()


def _add_options_table(
    document,
    headers: Sequence[str],
    options: Sequence[str],
    missing_value_text: str,
) -> int:
    table = document.add_table(rows=1, cols=len(headers))
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header

    missing_values = 0
    if not options:
        row_cells = table.add_row().cells
        row_cells[0].text = "1"
        if len(row_cells) > 1:
            row_cells[1].text = missing_value_text
            missing_values += 1
        for cell in row_cells[2:]:
            cell.text = missing_value_text
            missing_values += 1
        return missing_values

    for idx, option in enumerate(options, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        if len(row_cells) > 1:
            row_cells[1].text = option
        for cell in row_cells[2:]:
            cell.text = missing_value_text
            missing_values += 1
    return missing_values
